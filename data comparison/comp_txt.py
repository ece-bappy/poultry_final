import pandas as pd
import scipy.stats as stats
from docx import Document
from docx.shared import Inches

# Read the data from the CSV files
traditional_data = pd.read_csv("traditional_poultry_shed_data.csv")
controlled_data = pd.read_csv("controlled_poultry_shed_data.csv")


# Function to return descriptive statistics as a string
def get_descriptive_stats(data, label):
    return data.describe().to_string()  # Return as string


# Save the summary to a Microsoft Word document
doc = Document()
doc.add_heading('Analysis Summary', level=1)

# Add descriptive statistics tables
doc.add_heading('Descriptive statistics for Traditional Environment:', level=2)
traditional_table = doc.add_table(rows=1, cols=len(traditional_data.columns))
hdr_cells = traditional_table.rows[0].cells
for i, col in enumerate(traditional_data.columns):
    hdr_cells[i].text = col
for index, row in traditional_data.iterrows():
    row_cells = traditional_table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

doc.add_heading('Descriptive statistics for Controlled Environment:', level=2)
controlled_table = doc.add_table(rows=1, cols=len(controlled_data.columns))
hdr_cells = controlled_table.rows[0].cells
for i, col in enumerate(controlled_data.columns):
    hdr_cells[i].text = col
for index, row in controlled_data.iterrows():
    row_cells = controlled_table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# T-tests
parameters = ["Temperature", "Humidity", "Gas Level"]
t_test_results = []

for parameter in parameters:
    t_stat, p_value = stats.ttest_ind(
        traditional_data[parameter], controlled_data[parameter], equal_var=False
    )
    t_test_results.append([parameter, t_stat, p_value])

# Add t-test results in a table
doc.add_heading('T-test Results', level=2)
t_test_table = doc.add_table(rows=1, cols=3)
hdr_cells = t_test_table.rows[0].cells
hdr_cells[0].text = 'Parameter'
hdr_cells[1].text = 't-statistic'
hdr_cells[2].text = 'p-value'

for result in t_test_results:
    row_cells = t_test_table.add_row().cells
    row_cells[0].text = result[0]
    row_cells[1].text = str(result[1])
    row_cells[2].text = str(result[2])

# Save the document
doc.save('analysis_summary.docx')
